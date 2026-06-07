"""文档加载器 - 支持 PDF, Markdown, TXT, DOCX 格式."""

from pathlib import Path
from dataclasses import dataclass, field
from typing import Iterator
import re


@dataclass
class Document:
    """加载后的文档，保留元信息供后续引用."""

    content: str
    source: str
    page_num: int | None = None
    metadata: dict = field(default_factory=dict)


class DocumentLoader:
    """多格式文档加载器，返回统一的 Document 列表."""

    def load(self, file_path: str | Path) -> list[Document]:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {path}")

        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._load_pdf(path)
        if suffix in (".md", ".markdown", ".txt", ".py", ".json", ".yaml", ".yml"):
            return self._load_text(path)
        if suffix == ".docx":
            return self._load_docx(path)
        raise ValueError(f"不支持的文件格式: {suffix}")

    def load_directory(self, dir_path: str | Path, glob_pattern: str = "*") -> list[Document]:
        """批量加载目录下所有支持的文档."""
        docs: list[Document] = []
        supported = {".pdf", ".md", ".markdown", ".txt", ".py", ".json", ".yaml", ".yml", ".docx"}
        for fpath in Path(dir_path).rglob(glob_pattern):
            if fpath.suffix.lower() in supported:
                try:
                    docs.extend(self.load(fpath))
                except Exception:
                    pass
        return docs

    def _load_pdf(self, path: Path) -> list[Document]:
        from PyPDF2 import PdfReader

        reader = PdfReader(str(path))
        return [
            Document(
                content=page.extract_text() or "",
                source=str(path),
                page_num=i + 1,
                metadata={"total_pages": len(reader.pages)},
            )
            for i, page in enumerate(reader.pages)
        ]

    def _load_text(self, path: Path) -> list[Document]:
        content = path.read_text(encoding="utf-8", errors="replace")
        return [Document(content=content, source=str(path))]

    def _load_docx(self, path: Path) -> list[Document]:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # 也读取表格中的文字
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return [Document(content="\n".join(paragraphs), source=str(path))]
